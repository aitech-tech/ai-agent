"""
Tests for skills/word_skill_importer.py.

Creates temporary .docx files for both simple and advanced formats,
imports them, and verifies the generated JSON loads through SkillExecutor.
"""
import sys
import json
import tempfile
import shutil
from pathlib import Path

sys.path.insert(0, str(Path(__file__).parent.parent))


# ---------------------------------------------------------------------------
# Helpers — build minimal .docx files in memory
# ---------------------------------------------------------------------------

def _make_simple_docx(path: Path, overrides: dict = None) -> None:
    from docx import Document
    doc = Document()
    sections = {
        "Skill Name": "Overdue Invoice Review",
        "When I Ask": "show overdue invoices",
        "What This Skill Should Do": "Fetch all overdue invoices and group by status",
        "Use These Filters": "Status: overdue\nLimit: 30",
        "Use These Business Rules": "Flag invoices above 50000",
        "Approval Needed Before": "No",
        "Output Format": "Group by customer",
        "Notes": "Indian GST applies",
    }
    if overrides:
        sections.update(overrides)
    for title, value in sections.items():
        doc.add_heading(f"{title}:", level=2)
        doc.add_paragraph(value)
    doc.save(str(path))


def _make_advanced_docx(path: Path, rows=None) -> None:
    from docx import Document
    doc = Document()
    doc.add_paragraph("Skill Name: Advanced Invoice Skill")
    doc.add_paragraph("Description: Fetch unpaid invoices")
    doc.add_paragraph("Version: 1.0")
    doc.add_heading("Steps", level=1)
    table = doc.add_table(rows=1, cols=4)
    hdr = table.rows[0].cells
    hdr[0].text = "Step Name"
    hdr[1].text = "Tool"
    hdr[2].text = "Params"
    hdr[3].text = "On Error"
    if rows is None:
        rows = [("fetch", "zoho_books_list_invoices", '{"status":"unpaid"}', "stop")]
    for r in rows:
        cells = table.add_row().cells
        for i, v in enumerate(r):
            cells[i].text = v
    doc.save(str(path))


def _make_executor(skill: dict):
    """Load a skill dict into SkillExecutor and return executor."""
    from agent.skill_executor import SkillExecutor
    executor = SkillExecutor({})
    skill_id = f"zoho_books.{skill['name']}"
    executor._skills[skill_id] = skill
    return executor, skill_id


# ---------------------------------------------------------------------------
# A. Simple format
# ---------------------------------------------------------------------------

def test_simple_format_detected():
    """Simple template must be detected as 'simple' format."""
    from docx import Document
    from skills.word_skill_importer import _detect_format
    with tempfile.TemporaryDirectory() as tmp:
        p = Path(tmp) / "test.docx"
        _make_simple_docx(p)
        doc = Document(str(p))
        fmt = _detect_format(doc)
    assert fmt == "simple", f"Expected simple, got {fmt}"
    print("PASS: test_simple_format_detected")


def test_simple_import_generates_valid_json():
    """Importing a simple .docx must produce valid skill JSON with correct tool."""
    from skills.word_skill_importer import import_skill_from_word, CLIENT_SKILLS_DIR
    with tempfile.TemporaryDirectory() as tmp:
        p = Path(tmp) / "overdue_invoice_review.docx"
        _make_simple_docx(p)
        result = import_skill_from_word(str(p))

    assert result.get("success"), f"Import failed: {result.get('error')}"
    assert result["format"] == "simple"
    assert result["skill_id"].startswith("zoho_books.")
    assert result["steps"] >= 1

    # Read back and check tool
    skill = json.loads(Path(result["path"]).read_text(encoding="utf-8"))
    assert skill["steps"][0]["tool"] == "zoho_books_list_invoices"
    assert skill["steps"][0]["params"].get("status") == "overdue"
    print(f"PASS: test_simple_import_generates_valid_json (skill_id={result['skill_id']})")


def test_simple_import_sets_limit_from_filters():
    """Filter text '30 records' must set limit=30 in params."""
    from skills.word_skill_importer import import_skill_from_word
    with tempfile.TemporaryDirectory() as tmp:
        p = Path(tmp) / "limit_test.docx"
        _make_simple_docx(p, {"Skill Name": "Limit Test", "Use These Filters": "30 records"})
        result = import_skill_from_word(str(p))
    assert result.get("success"), result.get("error")
    skill = json.loads(Path(result["path"]).read_text(encoding="utf-8"))
    assert skill["steps"][0]["params"].get("limit") == 30
    print("PASS: test_simple_import_sets_limit_from_filters")


def test_simple_no_approval_for_read_only():
    """A read-only simple skill must not set approval_required."""
    from skills.word_skill_importer import import_skill_from_word
    with tempfile.TemporaryDirectory() as tmp:
        p = Path(tmp) / "readonly.docx"
        _make_simple_docx(p, {
            "Skill Name": "Read Only Expenses",
            "When I Ask": "show expenses",
            "What This Skill Should Do": "Fetch all expenses",
            "Approval Needed Before": "No",
        })
        result = import_skill_from_word(str(p))
    assert result.get("success"), result.get("error")
    assert result.get("approval_required") is False
    print("PASS: test_simple_no_approval_for_read_only")


def test_simple_approval_flagged_for_create():
    """A simple skill that mentions 'create' must set approval_required=True."""
    from skills.word_skill_importer import import_skill_from_word
    with tempfile.TemporaryDirectory() as tmp:
        p = Path(tmp) / "create_test.docx"
        _make_simple_docx(p, {
            "Skill Name": "Create Invoice Skill",
            "When I Ask": "create a new invoice",
            "What This Skill Should Do": "Create an invoice for the customer",
            "Approval Needed Before": "Yes — creating records needs approval",
        })
        result = import_skill_from_word(str(p))
    assert result.get("success"), result.get("error")
    assert result.get("approval_required") is True, "approval_required should be True for create action"
    print("PASS: test_simple_approval_flagged_for_create")


def test_simple_transform_step_added_when_grouping_requested():
    """Output format mentioning 'group' must add a summarise step."""
    from skills.word_skill_importer import import_skill_from_word
    with tempfile.TemporaryDirectory() as tmp:
        p = Path(tmp) / "group_test.docx"
        _make_simple_docx(p, {
            "Skill Name": "Group By Status",
            "When I Ask": "invoice summary grouped by status",
            "What This Skill Should Do": "Summarise invoices grouped by status",
            "Output Format": "Group by status and show totals",
        })
        result = import_skill_from_word(str(p))
    assert result.get("success"), result.get("error")
    skill = json.loads(Path(result["path"]).read_text(encoding="utf-8"))
    step_tools = [s["tool"] for s in skill["steps"]]
    assert "summarise_results" in step_tools, f"Expected summarise_results step, got: {step_tools}"
    print("PASS: test_simple_transform_step_added_when_grouping_requested")


def test_simple_customer_task_maps_correctly():
    """'show customers' must map to zoho_books_list_contacts with contact_type=customer."""
    from skills.word_skill_importer import import_skill_from_word
    with tempfile.TemporaryDirectory() as tmp:
        p = Path(tmp) / "customer_review.docx"
        _make_simple_docx(p, {
            "Skill Name": "Customer Review",
            "When I Ask": "show customers",
            "What This Skill Should Do": "Fetch all customers",
            "Approval Needed Before": "No",
        })
        result = import_skill_from_word(str(p))
    assert result.get("success"), result.get("error")
    skill = json.loads(Path(result["path"]).read_text(encoding="utf-8"))
    step = skill["steps"][0]
    assert step["tool"] == "zoho_books_list_contacts"
    assert step["params"].get("contact_type") == "customer"
    print("PASS: test_simple_customer_task_maps_correctly")


def test_simple_skill_loads_in_skill_executor():
    """Generated simple skill JSON must load in SkillExecutor without errors."""
    from skills.word_skill_importer import import_skill_from_word
    with tempfile.TemporaryDirectory() as tmp:
        p = Path(tmp) / "executor_test.docx"
        _make_simple_docx(p, {"Skill Name": "Executor Test Skill"})
        result = import_skill_from_word(str(p))
    assert result.get("success"), result.get("error")
    skill = json.loads(Path(result["path"]).read_text(encoding="utf-8"))
    executor, skill_id = _make_executor(skill)
    found = next((s for s in executor.list_skills() if s["id"] == skill_id), None)
    assert found is not None, f"Skill {skill_id} not found in SkillExecutor"
    print(f"PASS: test_simple_skill_loads_in_skill_executor ({skill_id})")


# ---------------------------------------------------------------------------
# B. Advanced format
# ---------------------------------------------------------------------------

def test_advanced_format_detected():
    """Advanced template (with Tool table) must be detected as 'advanced'."""
    from docx import Document
    from skills.word_skill_importer import _detect_format
    with tempfile.TemporaryDirectory() as tmp:
        p = Path(tmp) / "adv.docx"
        _make_advanced_docx(p)
        doc = Document(str(p))
        fmt = _detect_format(doc)
    assert fmt == "advanced", f"Expected advanced, got {fmt}"
    print("PASS: test_advanced_format_detected")


def test_advanced_import_generates_valid_json():
    """Advanced .docx import must produce valid skill JSON with correct steps."""
    from skills.word_skill_importer import import_skill_from_word
    with tempfile.TemporaryDirectory() as tmp:
        p = Path(tmp) / "advanced_invoice_skill.docx"
        _make_advanced_docx(p)
        result = import_skill_from_word(str(p))
    assert result.get("success"), f"Import failed: {result.get('error')}"
    assert result["format"] == "advanced"
    skill = json.loads(Path(result["path"]).read_text(encoding="utf-8"))
    assert skill["steps"][0]["tool"] == "zoho_books_list_invoices"
    assert skill["steps"][0]["params"].get("status") == "unpaid"
    print(f"PASS: test_advanced_import_generates_valid_json (skill_id={result['skill_id']})")


def test_advanced_delete_tool_blocked():
    """Delete tools in an advanced skill table must be silently removed."""
    from skills.word_skill_importer import import_skill_from_word
    with tempfile.TemporaryDirectory() as tmp:
        p = Path(tmp) / "delete_test.docx"
        _make_advanced_docx(p, rows=[
            ("fetch", "zoho_books_list_invoices", "{}", "stop"),
            ("delete", "zoho_books_delete_invoice", '{"invoice_id":"123"}', "stop"),
        ])
        result = import_skill_from_word(str(p))
    # Delete step should be stripped → skill still valid with only fetch step
    assert result.get("success"), f"Expected success after stripping delete step: {result}"
    skill = json.loads(Path(result["path"]).read_text(encoding="utf-8"))
    step_tools = [s["tool"] for s in skill["steps"]]
    assert "zoho_books_delete_invoice" not in step_tools, "Delete tool must be blocked"
    assert "zoho_books_list_invoices" in step_tools
    print("PASS: test_advanced_delete_tool_blocked")


def test_advanced_skill_loads_in_skill_executor():
    """Generated advanced skill JSON must load in SkillExecutor without errors."""
    from skills.word_skill_importer import import_skill_from_word
    with tempfile.TemporaryDirectory() as tmp:
        p = Path(tmp) / "adv_executor_test.docx"
        _make_advanced_docx(p)
        result = import_skill_from_word(str(p))
    assert result.get("success"), result.get("error")
    skill = json.loads(Path(result["path"]).read_text(encoding="utf-8"))
    executor, skill_id = _make_executor(skill)
    found = next((s for s in executor.list_skills() if s["id"] == skill_id), None)
    assert found is not None, f"Skill {skill_id} not found in SkillExecutor"
    print(f"PASS: test_advanced_skill_loads_in_skill_executor ({skill_id})")


# ---------------------------------------------------------------------------
# C. Filename-only resolution
# ---------------------------------------------------------------------------

def test_import_by_filename_resolves_from_templates_dir():
    """import_skill_from_word must accept just a filename and find it in TEMPLATES_DIR."""
    from skills.word_skill_importer import import_skill_from_word, TEMPLATES_DIR
    TEMPLATES_DIR.mkdir(parents=True, exist_ok=True)
    fname = "_test_filename_resolution.docx"
    tmp_path = TEMPLATES_DIR / fname
    try:
        _make_simple_docx(tmp_path, {"Skill Name": "Filename Resolution Test"})
        result = import_skill_from_word(fname)
        assert result.get("success"), f"Filename resolution failed: {result.get('error')}"
        print(f"PASS: test_import_by_filename_resolves_from_templates_dir")
    finally:
        tmp_path.unlink(missing_ok=True)
        # Clean up generated JSON
        from skills.word_skill_importer import CLIENT_SKILLS_DIR
        (CLIENT_SKILLS_DIR / "filename_resolution_test.json").unlink(missing_ok=True)


# ---------------------------------------------------------------------------
# D. Validate tool
# ---------------------------------------------------------------------------

def test_validate_client_skill_passes_for_valid():
    """validate_client_skill must return valid=True for a well-formed skill."""
    from skills.word_skill_importer import import_skill_from_word, validate_skill
    with tempfile.TemporaryDirectory() as tmp:
        p = Path(tmp) / "validate_me.docx"
        _make_simple_docx(p, {"Skill Name": "Validate Me"})
        result = import_skill_from_word(str(p))
    assert result.get("success"), result.get("error")
    skill = json.loads(Path(result["path"]).read_text(encoding="utf-8"))
    errors = validate_skill(skill)
    assert not errors, f"Unexpected validation errors: {errors}"
    print("PASS: test_validate_client_skill_passes_for_valid")


def test_validate_rejects_missing_name():
    """validate_skill must reject a skill dict missing 'name'."""
    from skills.word_skill_importer import validate_skill
    errors = validate_skill({"steps": []})
    assert any("name" in e for e in errors), f"Expected missing name error, got: {errors}"
    print("PASS: test_validate_rejects_missing_name")


def test_validate_rejects_blocked_delete_tool():
    """validate_skill must reject a step using a delete tool."""
    from skills.word_skill_importer import validate_skill
    skill = {
        "name": "bad_skill",
        "steps": [{"step_name": "s", "tool": "zoho_books_delete_invoice", "params": {}, "on_error": "stop"}],
    }
    errors = validate_skill(skill)
    assert any("delete" in e.lower() or "blocked" in e.lower() for e in errors), f"Expected blocked tool error: {errors}"
    print("PASS: test_validate_rejects_blocked_delete_tool")


# ---------------------------------------------------------------------------
# E. MCP tools smoke test
# ---------------------------------------------------------------------------

def test_mcp_list_skill_templates_returns_dict():
    """list_skill_templates MCP tool must return success:true with templates list."""
    from tools.word_skill_tools import _tool_list_skill_templates
    result = _tool_list_skill_templates({})
    assert result.get("success"), result
    assert "templates" in result
    assert isinstance(result["templates"], list)
    print(f"PASS: test_mcp_list_skill_templates_returns_dict ({result['count']} templates)")


def test_mcp_list_client_skills_returns_dict():
    """list_client_skills MCP tool must return success:true with client_skills list."""
    from tools.word_skill_tools import _tool_list_client_skills
    result = _tool_list_client_skills({})
    assert result.get("success"), result
    assert "client_skills" in result
    print(f"PASS: test_mcp_list_client_skills_returns_dict ({result['count']} skills)")


def test_mcp_import_missing_file_returns_error():
    """import_skill_from_word with a non-existent file must return success:false."""
    from tools.word_skill_tools import _tool_import_skill_from_word
    result = _tool_import_skill_from_word({"path": "nonexistent_file_xyz.docx"})
    assert result.get("success") is False
    assert "error" in result
    print("PASS: test_mcp_import_missing_file_returns_error")


if __name__ == "__main__":
    # Simple format
    test_simple_format_detected()
    test_simple_import_generates_valid_json()
    test_simple_import_sets_limit_from_filters()
    test_simple_no_approval_for_read_only()
    test_simple_approval_flagged_for_create()
    test_simple_transform_step_added_when_grouping_requested()
    test_simple_customer_task_maps_correctly()
    test_simple_skill_loads_in_skill_executor()
    # Advanced format
    test_advanced_format_detected()
    test_advanced_import_generates_valid_json()
    test_advanced_delete_tool_blocked()
    test_advanced_skill_loads_in_skill_executor()
    # Filename resolution
    test_import_by_filename_resolves_from_templates_dir()
    # Validation
    test_validate_client_skill_passes_for_valid()
    test_validate_rejects_missing_name()
    test_validate_rejects_blocked_delete_tool()
    # MCP tools
    test_mcp_list_skill_templates_returns_dict()
    test_mcp_list_client_skills_returns_dict()
    test_mcp_import_missing_file_returns_error()
    print("\nAll Word skill importer tests passed.")
