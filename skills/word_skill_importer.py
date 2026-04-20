"""
word_skill_importer.py — ReckLabs AI Agent

Converts .docx skill templates into validated skill JSON.

Supports two formats auto-detected from document content:
  - Simple  : business-language sections, no tool knowledge needed
  - Advanced: table-based format for power users (original behaviour)

Generated JSON is saved to skills/client/zoho_books/<skill_id>.json
and loads as zoho_books.<skill_id> in SkillExecutor.
"""
import json
import re
import logging
from pathlib import Path

logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# Paths
# ---------------------------------------------------------------------------

_BASE = Path(__file__).parent.parent
TEMPLATES_DIR = _BASE / "skills" / "client_docs" / "zoho_books"
CLIENT_SKILLS_DIR = _BASE / "skills" / "client" / "zoho_books"

# ---------------------------------------------------------------------------
# Keyword → tool mapping for simple format
# ---------------------------------------------------------------------------

_TASK_MAP = [
    (r"overdue\s+invoice", "zoho_books_list_invoices", {"status": "overdue"}),
    (r"unpaid\s+invoice", "zoho_books_list_invoices", {"status": "unpaid"}),
    (r"draft\s+invoice", "zoho_books_list_invoices", {"status": "draft"}),
    (r"invoice", "zoho_books_list_invoices", {}),
    (r"overdue\s+(bill|payable)", "zoho_books_list_expenses", {"status": "overdue"}),
    (r"expense|bill|payable", "zoho_books_list_expenses", {}),
    (r"customer", "zoho_books_list_contacts", {"contact_type": "customer"}),
    (r"vendor|supplier", "zoho_books_list_contacts", {"contact_type": "vendor"}),
    (r"contact", "zoho_books_list_contacts", {}),
    (r"item|product|service", "zoho_books_list_items", {}),
    (r"tax|gst|tds", "zoho_books_list_taxes", {}),
    (r"estimate|quote", "zoho_books_list_estimates", {}),
    (r"sales.?order", "zoho_books_list_sales_orders", {}),
    (r"purchase.?order|\bpo\b", "zoho_books_list_purchase_orders", {}),
    (r"payment|receipt", "zoho_books_list_customer_payments", {}),
]

# Tools that must never appear in a simple-format skill
_BLOCKED_TOOLS = {
    "zoho_books_delete_invoice", "zoho_books_delete_contact",
    "zoho_books_delete_expense", "zoho_books_delete_item",
    "zoho_books_delete_tax", "zoho_books_delete_estimate",
    "zoho_books_delete_sales_order", "zoho_books_delete_purchase_order",
    "zoho_books_delete_customer_payment",
}

# Actions that require human approval before executing
_APPROVAL_KEYWORDS = re.compile(
    r"\b(create|new|add|send|email|update|edit|change|record|pay|mark)\b", re.I
)

# Summarise/group transform triggers
_TRANSFORM_KEYWORDS = re.compile(
    r"\b(summar|group|breakdown|by status|by customer|by vendor|"
    r"by category|by contact|total|aggregate|count)\b", re.I
)

# ---------------------------------------------------------------------------
# Simple format: known section headings
# ---------------------------------------------------------------------------

_SIMPLE_SECTIONS = [
    "skill name",
    "when i ask",
    "what this skill should do",
    "use these filters",
    "use these business rules",
    "approval needed before",
    "output format",
    "notes",
]

# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _slugify(text: str) -> str:
    text = text.lower().strip()
    text = re.sub(r"[^a-z0-9]+", "_", text)
    return text.strip("_")[:60]


def _detect_format(doc) -> str:
    """Return 'simple' or 'advanced' based on document content."""
    full_text = "\n".join(p.text.strip() for p in doc.paragraphs if p.text.strip())
    lower = full_text.lower()

    # Advanced: has a table with a Tool column
    for table in doc.tables:
        if not table.rows:
            continue
        headers = [c.text.strip().lower() for c in table.rows[0].cells]
        if "tool" in headers:
            return "advanced"

    # Simple: matches at least 2 known section headings
    hits = sum(1 for s in _SIMPLE_SECTIONS if s in lower)
    if hits >= 2:
        return "simple"

    return "advanced"


def _extract_sections(doc) -> dict:
    """Extract key→value sections from a simple-format document."""
    sections: dict = {}
    current_key = None
    current_lines: list = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        lower = text.lower().rstrip(":")
        matched = None
        for s in _SIMPLE_SECTIONS:
            if lower == s or lower.startswith(s):
                matched = s
                break

        if matched:
            if current_key and current_lines:
                sections[current_key] = " ".join(current_lines).strip()
            current_key = matched
            current_lines = []
            if ":" in text:
                after = text.split(":", 1)[1].strip()
                if after:
                    current_lines.append(after)
        elif current_key is not None:
            current_lines.append(text)

    if current_key and current_lines:
        sections[current_key] = " ".join(current_lines).strip()

    return sections


def _resolve_tool(task_text: str) -> tuple:
    """Match task description to a Zoho Books list tool + base params."""
    lower = task_text.lower()
    for pattern, tool, params in _TASK_MAP:
        if re.search(pattern, lower):
            return tool, dict(params)
    return "zoho_books_list_invoices", {}


def _parse_filters(filter_text: str, base_params: dict) -> dict:
    """Extract simple filter values from free text and merge into params."""
    params = dict(base_params)
    lower = filter_text.lower()

    m = re.search(r"\b(\d+)\s*(?:records?|items?|rows?|results?)?", lower)
    if m:
        params["limit"] = int(m.group(1))

    for status in ("overdue", "unpaid", "paid", "draft", "sent", "accepted", "declined"):
        if status in lower:
            params.setdefault("status", status)
            break

    return params


def _compile_simple(sections: dict, source_file: str) -> dict:
    """Convert simple-format sections into a validated skill dict."""
    raw_name = sections.get("skill name", "").strip() or Path(source_file).stem
    skill_id = _slugify(raw_name)
    description = (
        sections.get("what this skill should do", "")
        or sections.get("when i ask", "")
        or raw_name
    )

    task_text = " ".join([
        sections.get("when i ask", ""),
        sections.get("what this skill should do", ""),
    ])

    tool_name, base_params = _resolve_tool(task_text)
    filter_text = sections.get("use these filters", "")
    params = _parse_filters(filter_text, base_params) if filter_text else dict(base_params)

    # Approval: explicit field overrides keyword scan
    approval_field = sections.get("approval needed before", "")
    if re.search(r"\b(no|none|not required|n/a)\b", approval_field, re.I):
        needs_approval = False
    elif approval_field and approval_field.strip():
        needs_approval = True
    else:
        needs_approval = bool(_APPROVAL_KEYWORDS.search(task_text))

    fetch_step: dict = {
        "step_name": "fetch_data",
        "tool": tool_name,
        "params": params,
        "on_error": "stop",
    }
    if needs_approval:
        fetch_step["approval_required"] = True

    steps = [fetch_step]

    # Auto-add summarise transform when user asks for grouping/totals
    if _TRANSFORM_KEYWORDS.search(task_text + " " + sections.get("output format", "")):
        steps.append({
            "step_name": "summarise",
            "tool": "summarise_results",
            "params": {"group_by": "status", "source_step": "fetch_data"},
            "on_error": "continue",
        })

    skill: dict = {
        "name": skill_id,
        "display_name": raw_name,
        "description": description,
        "version": "1.0",
        "format": "simple",
        "source_doc": Path(source_file).name,
        "steps": steps,
    }
    if sections.get("use these business rules"):
        skill["business_rules"] = sections["use these business rules"]
    if sections.get("notes"):
        skill["notes"] = sections["notes"]
    if needs_approval:
        skill["approval_required"] = True

    return skill


# ---------------------------------------------------------------------------
# Advanced format (table-based)
# ---------------------------------------------------------------------------

def _parse_advanced(doc, source_file: str) -> dict:
    """Parse the table-based advanced format."""
    meta: dict = {}
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        if ":" in text:
            key, _, val = text.partition(":")
            key_low = key.strip().lower()
            val = val.strip()
            if key_low in ("skill name", "name"):
                meta["name"] = val
            elif key_low == "description":
                meta["description"] = val
            elif key_low == "version":
                meta["version"] = val

    skill_id = _slugify(meta.get("name", Path(source_file).stem))
    steps = []

    for table in doc.tables:
        if not table.rows:
            continue
        headers = [c.text.strip().lower() for c in table.rows[0].cells]
        if "tool" not in headers:
            continue

        col = {h: i for i, h in enumerate(headers)}

        for row in table.rows[1:]:
            cells = [c.text.strip() for c in row.cells]
            if not any(cells):
                continue

            tool = cells[col["tool"]] if col.get("tool") is not None else ""
            if not tool:
                continue
            if tool in _BLOCKED_TOOLS:
                logger.warning("Blocked delete tool '%s' — step removed from advanced skill", tool)
                continue

            # Step name
            sn_key = next((k for k in col if "step" in k), None)
            step_name = cells[col[sn_key]] if sn_key and col[sn_key] < len(cells) else f"step_{len(steps)+1}"

            # on_error
            oe_key = next((k for k in col if "error" in k), None)
            on_error = cells[col[oe_key]].lower() if oe_key and col[oe_key] < len(cells) and cells[col[oe_key]] else "stop"

            # params
            params: dict = {}
            p_key = next((k for k in col if "param" in k), None)
            if p_key and col[p_key] < len(cells):
                params_raw = cells[col[p_key]]
                if params_raw:
                    try:
                        params = json.loads(params_raw)
                    except json.JSONDecodeError:
                        for line in re.split(r"[\n,;]+", params_raw):
                            if "=" in line:
                                k2, _, v2 = line.partition("=")
                                params[k2.strip()] = v2.strip()

            steps.append({
                "step_name": step_name,
                "tool": tool,
                "params": params,
                "on_error": on_error if on_error in ("stop", "continue") else "stop",
            })
        break  # first matching table only

    return {
        "name": skill_id,
        "display_name": meta.get("name", skill_id),
        "description": meta.get("description", ""),
        "version": meta.get("version", "1.0"),
        "format": "advanced",
        "source_doc": Path(source_file).name,
        "steps": steps,
    }


# ---------------------------------------------------------------------------
# Validation
# ---------------------------------------------------------------------------

_REQUIRED_FIELDS = {"name", "steps"}
_VALID_ON_ERROR = {"stop", "continue"}


def validate_skill(skill: dict) -> list:
    """Return list of validation error strings (empty list = valid)."""
    errors = []
    for f in _REQUIRED_FIELDS:
        if f not in skill:
            errors.append(f"Missing required field: '{f}'")
    if "name" in skill and not re.match(r"^[a-z0-9_]{1,60}$", skill["name"]):
        errors.append(f"'name' must be lowercase a-z/0-9/_ only, got: {skill['name']!r}")
    for i, step in enumerate(skill.get("steps", [])):
        if "tool" not in step:
            errors.append(f"Step {i+1} missing 'tool'")
        if step.get("on_error", "stop") not in _VALID_ON_ERROR:
            errors.append(f"Step {i+1} invalid on_error: {step.get('on_error')!r}")
        if step.get("tool") in _BLOCKED_TOOLS:
            errors.append(f"Step {i+1} uses blocked delete tool: {step['tool']!r}")
    return errors


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------

def import_skill_from_word(path_or_filename: str) -> dict:
    """
    Import a .docx skill template, convert to JSON, save to client skills dir.

    Accepts a full path OR just a filename from skills/client_docs/zoho_books/.

    Returns success dict or {"success": False, "error": "..."}.
    """
    try:
        from docx import Document
    except ImportError:
        return {"success": False, "error": "python-docx not installed — run: pip install python-docx"}

    # Resolve path: try as-is, then look in TEMPLATES_DIR
    p = Path(path_or_filename)
    if not p.exists():
        for candidate in [
            TEMPLATES_DIR / path_or_filename,
            TEMPLATES_DIR / (path_or_filename + ".docx"),
        ]:
            if candidate.exists():
                p = candidate
                break

    if not p.exists():
        return {"success": False, "error": f"File not found: {path_or_filename}"}

    try:
        doc = Document(str(p))
    except Exception as e:
        return {"success": False, "error": f"Could not open Word document: {e}"}

    fmt = _detect_format(doc)
    warnings = []

    try:
        if fmt == "simple":
            sections = _extract_sections(doc)
            if not sections:
                warnings.append("No recognised sections found — falling back to advanced parser")
                fmt = "advanced"
                skill = _parse_advanced(doc, str(p))
            else:
                skill = _compile_simple(sections, str(p))
        else:
            skill = _parse_advanced(doc, str(p))
    except Exception as e:
        return {"success": False, "error": f"Parse error ({fmt} format): {e}"}

    errors = validate_skill(skill)
    if errors:
        return {"success": False, "error": "Validation failed: " + "; ".join(errors)}

    CLIENT_SKILLS_DIR.mkdir(parents=True, exist_ok=True)
    out_path = CLIENT_SKILLS_DIR / f"{skill['name']}.json"
    out_path.write_text(json.dumps(skill, indent=2, ensure_ascii=False), encoding="utf-8")
    logger.info("Imported skill '%s' (%s format) from %s → %s", skill["name"], fmt, p.name, out_path)

    return {
        "success": True,
        "skill_id": f"zoho_books.{skill['name']}",
        "path": str(out_path),
        "format": fmt,
        "display_name": skill.get("display_name", skill["name"]),
        "steps": len(skill.get("steps", [])),
        "approval_required": skill.get("approval_required", False),
        "warnings": warnings,
    }


def list_skill_templates() -> list:
    """Return user-friendly list of Word templates in TEMPLATES_DIR."""
    TEMPLATES_DIR.mkdir(parents=True, exist_ok=True)
    results = []
    for f in sorted(TEMPLATES_DIR.glob("*.docx")):
        friendly = f.stem.replace("_", " ").replace("-", " ").title()
        results.append({
            "filename": f.name,
            "display_name": friendly,
            "path": str(f),
            "size_kb": round(f.stat().st_size / 1024, 1),
        })
    return results


def list_client_skills() -> list:
    """Return list of generated client skill JSON files."""
    CLIENT_SKILLS_DIR.mkdir(parents=True, exist_ok=True)
    results = []
    for f in sorted(CLIENT_SKILLS_DIR.glob("*.json")):
        try:
            skill = json.loads(f.read_text(encoding="utf-8"))
            results.append({
                "skill_id": f"zoho_books.{skill.get('name', f.stem)}",
                "display_name": skill.get("display_name", skill.get("name", f.stem)),
                "description": skill.get("description", ""),
                "format": skill.get("format", "unknown"),
                "source_doc": skill.get("source_doc", ""),
                "steps": len(skill.get("steps", [])),
                "approval_required": skill.get("approval_required", False),
                "path": str(f),
            })
        except Exception:
            results.append({
                "skill_id": f"zoho_books.{f.stem}",
                "path": str(f),
                "error": "invalid JSON",
            })
    return results
